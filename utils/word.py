import csv
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image

# ------------------ CONFIG ------------------
BASE_DIR = os.path.dirname(os.path.abspath(__file__))   # /utils
ROOT_DIR = os.path.abspath(os.path.join(BASE_DIR, ".."))  # repo root

CSV_FILE = os.path.join(ROOT_DIR, "TestData", "data.csv")
STEP_REPORT_FILE = os.path.join(ROOT_DIR, "test-output", "Automation_Step_Report.docx")
SCREENSHOT_REPORT_FILE = os.path.join(ROOT_DIR, "test-output", "Automation_Screenshot_Report.docx")
SCREENSHOT_DIR = os.path.join(ROOT_DIR, "test-output", "screenshots")

TAG = "Regression"
ENVIRONMENT = "QA"
URL = "http://example.com"  

os.makedirs(SCREENSHOT_DIR, exist_ok=True)

# ------------------ STYLING ------------------
def style_cell(cell, bg_color=None, bold=False):
    tc_pr = cell._tc.get_or_add_tcPr()
    if bg_color:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), bg_color)
        tc_pr.append(shd)
    if bold:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

def add_header_table(doc, run_time):
    """
    Header table with:
    Row1: Tag | Environment
    Row2: Values
    Row3: Date & Time | URL
    """
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"

    # Header row
    style_cell(table.rows[0].cells[0], "BDD7EE", True)
    style_cell(table.rows[0].cells[1], "BDD7EE", True)
    table.rows[0].cells[0].text = "Tag"
    table.rows[0].cells[1].text = "Environment"

    # Values row
    table.rows[1].cells[0].text = TAG
    table.rows[1].cells[1].text = ENVIRONMENT

    # Date & Time and URL row
    table.rows[2].cells[0].text = f"Date & Time: {run_time}"
    table.rows[2].cells[1].text = f"URL: {URL}"

    doc.add_paragraph("")

# ------------------ STEP REPORT ------------------
def create_step_report(rows):
    run_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    doc = Document()
    doc.add_heading("AUTOMATION EXECUTION REPORT - STEPS", 0)
    add_header_table(doc, run_time)

    ROWS_PER_PAGE = 12
    count = 0

    def new_table():
        t = doc.add_table(rows=1, cols=4)
        t.style = "Light Grid Accent 1"
        hdr = t.rows[0].cells
        hdr[0].text = "Step"
        hdr[1].text = "Description"
        hdr[2].text = "Expected"
        hdr[3].text = "Screenshot ID"
        for cell in hdr:
            style_cell(cell, "BDD7EE", True)
        return t

    table = new_table()

    for row in rows:
        if count > 0 and count % ROWS_PER_PAGE == 0:
            doc.add_page_break()
            add_header_table(doc, run_time)  # add header only
            table = new_table()  # start new table

        c = table.add_row().cells
        c[0].text = row["Step"]
        c[1].text = row["Description"]
        c[2].text = row["Expected"]
        c[3].text = row["ScreenshotID"]

        count += 1

    doc.save(STEP_REPORT_FILE)
    print(f"STEP REPORT CREATED: {STEP_REPORT_FILE}")

# ------------------ SCREENSHOT REPORT (1 screenshot per page, clickable ID) ------------------
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):
    """
    Add a clickable hyperlink to a paragraph in Word.
    """
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Color
    c = OxmlElement('w:color')
    c.set(qn('w:val'), color)
    rPr.append(c)

    # Underline
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    new_run.append(rPr)

    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def create_screenshot_report(rows):
    doc = Document()

    for row in rows:
        run_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        add_header_table(doc, run_time)

        sid = row["ScreenshotID"]
        step = row["Step"]
        img_path = os.path.join(SCREENSHOT_DIR, f"{sid}.png")

        # Paragraph with clickable Screenshot ID
        p = doc.add_paragraph(f"Step: {step} | Screenshot ID: ")
        if os.path.exists(img_path):
            add_hyperlink(p, img_path, sid)

        # Add image
        if os.path.exists(img_path):
            try:
                doc.add_picture(img_path, width=Inches(6))
            except Exception as e:
                print(f"Image insert failed: {e}")

        # Page break after each screenshot
        doc.add_page_break()

    doc.save(SCREENSHOT_REPORT_FILE)
    print(f"SCREENSHOT REPORT CREATED: {SCREENSHOT_REPORT_FILE}")


def main():
    if not os.path.exists(CSV_FILE):
        print("CSV not found")
        return

    with open(CSV_FILE, newline="", encoding="utf-8") as f:
        rows = list(csv.DictReader(f))

    if not rows:
        print("CSV is empty")
        return

    create_step_report(rows)
    create_screenshot_report(rows)

if __name__ == "__main__":
    main()
